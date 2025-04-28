from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import time
import docx
from docx import Document as DocxReader
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from openai import OpenAI, OpenAIError
import chromadb
from chromadb.config import Settings
import logging

# Initialize ChromaDB client
CHROMA_PERSIST_DIR = ".chromadb"

chroma_client = chromadb.PersistentClient(
    path=CHROMA_PERSIST_DIR
)

# get_or_create so re‑starts don’t duplicate documents
try:
    chroma_client.delete_collection("sdlc-rag-index")
except Exception:
    pass

# Create with cosine distance and a higher construction ef for better index quality
collection = chroma_client.create_collection(
    name="sdlc-rag-index",
    metadata={
        "hnsw:space": "cosine",            # use cosine distance instead of L2 :contentReference[oaicite:0]{index=0}
        "hnsw:construction_ef": 200,       # build a richer graph :contentReference[oaicite:1]{index=1}
        "hnsw:sync_threshold": 1000        # sync threshold (optional tweak)
    }
)

# Initialize OpenAI client (using environment variable or directly here if preferred)
client = OpenAI()

# 📁 Folder setup
DOCUMENT_TO_ANALYZE_PATH = "document_to_analyze"
PROPRIETARY_FOLDER = "proprietary_documents"
RESULTS_FOLDER = "results"

# Ensure required folders exist
os.makedirs(DOCUMENT_TO_ANALYZE_PATH, exist_ok=True)
os.makedirs(PROPRIETARY_FOLDER, exist_ok=True)
os.makedirs(RESULTS_FOLDER, exist_ok=True)


# 📄 Read .docx into text
def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def parse_trace_matrix(file_path: str):
    """
    Read the .docx and yield (req_id, req_text) pairs by paragraph.
    Starts a new entry whenever a paragraph begins with a valid ID,
    then accumulates all following paragraphs until the next ID.
    """
    doc = DocxReader(file_path)
    current_id = None
    current_text = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        # ID line?
        m = re.match(r'^(BR\s*\d+(?:\.\d+)*|FR\s*\d+(?:\.\d+)*|UR-REG-\d+|FS-REG-\d+)\b', txt)
        if m:
            # emit the last one
            if current_id:
                yield current_id, " ".join(current_text).strip()
            # start new
            current_id = m.group(1)
            rest = txt[len(current_id):].strip()
            current_text = [rest] if rest else []
        else:
            # continuation paragraph
            if current_id and txt:
                current_text.append(txt)

    # emit final entry
    if current_id:
        yield current_id, " ".join(current_text).strip()

def get_embedding(text: str) -> list[float]:
    """
    Call OpenAI to embed a single string, 
    returning the embedding vector.
    """
    response = client.embeddings.create(
        model="text-embedding-ada-002",
        input=[text]
    )
    # grab the first (and only) embedding
    return response.data[0].embedding

def retrieve_relevant_trace_requirements(change_description, top_k=5):
    emb = get_embedding(change_description)

    # 1) grab more candidates
    results = collection.query(
      query_embeddings=[emb],
      n_results=10,                # ← bump this up
      include=["metadatas","distances"]
    )

    # 2) pair them up
    candidates = list(zip(
      results["metadatas"][0],
      results["distances"][0]
    ))

    # 3) filter out anything too far (distance > .4, say)
    candidates = [
      (m["requirement_id"], m["text"], dist)
      for m, dist in candidates
      if dist < 0.4
    ]

    # 4) sort by ascending distance (i.e. best match first)
    candidates.sort(key=lambda x: x[2])

    # 5) take your top_k
    top = candidates[:top_k]

    # 6) format as before
    return [
      (req_id, text)
      for req_id, text, _ in top
    ]

# after you’ve done your imports, client & collection setup, and before app.run():
def ingest_trace_matrix():
    before = collection.count()
    app.logger.debug("Trace Matrix before ingest: %d", before)

    all_ids = []
    if before == 0:
        for file in os.listdir(PROPRIETARY_FOLDER):
            if file.lower().endswith(".docx") and "trace matrix v8" in file.lower():
                full_path = os.path.join(PROPRIETARY_FOLDER, file)
                for req_id, req_text in parse_trace_matrix(full_path):
                    if not req_text.strip():
                        continue
                    emb = get_embedding(req_text)
                    collection.add(
                        ids=[req_id],
                        metadatas=[{"requirement_id": req_id, "text": req_text}],
                        embeddings=[emb]
                    )

    after = collection.count()
    app.logger.debug("Trace Matrix after  ingest: %d", after)
    app.logger.debug("IDs ingested: %s", all_ids)

app = Flask(__name__)
app.logger.setLevel(logging.DEBUG)
CORS(app, origins=["https://sdlc-reviewer.azurewebsites.net/"])

# 💾 Save plain text to .docx
def save_to_docx(text, output_path):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(output_path)

# 📄 Format AI output into styled .docx
def format_ai_output(ai_text, output_path):
    doc = Document()
    lines = ai_text.strip().split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line.lstrip("# ").strip())

        if line.startswith("# "):
            run.bold = True
            run.font.size = Pt(22)
        elif line.startswith("## "):
            run.bold = True
            run.font.size = Pt(16)
        elif line.startswith("### "):
            run.bold = True
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(output_path)

# 🔍 Extract operational testing section using regex
def extract_operational_testing_section(text):
    match = re.search(
        r'performed to validate the following Medium and low risk functional requirements:(.*?)Purpose and Scope',
        text, re.DOTALL
    )
    if match:
        return match.group(1).strip()
    return None

# 🤖 Run prompt using Assistant API
def run_prompt_with_assistant(assistant_id, user_input):
    """Runs a prompt using the Assistant API, returns response and metrics."""
    try:
        start_time = time.time()

        thread = client.beta.threads.create()
        client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=user_input
        )

        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id,
        )

        while True:
            run_status = client.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id
            )
            if run_status.status == "completed":
                break
            elif run_status.status == "failed":
                print(f"❌ Assistant failed with error: {run_status.last_error}")
                raise Exception(f"Assistant run failed: {run_status.last_error}")
            elif run_status.status in ["cancelled", "expired"]:
                raise Exception(f"Assistant run failed: {run_status.status}")
            time.sleep(1)

        # ⏱️ Elapsed time
        elapsed_time = time.time() - start_time

        # 📊 Messages + Token usage (requires separate API call)
        messages = client.beta.threads.messages.list(thread_id=thread.id)
        result_text = messages.data[0].content[0].text.value

        # 🔢 Token count and cost (Estimate based on model pricing — adjust as needed)
        usage = run_status.usage  # Only available if `retrieval_tool` or `code_interpreter` not enabled
        if usage:
            tokens_used = usage.total_tokens
            cost = tokens_used / 1000 * 0.005  # Example: $0.005 / 1k for gpt-4o
        else:
            tokens_used = 0
            cost = 0.0

        return result_text, tokens_used, cost, elapsed_time

    except Exception as e:
        raise Exception(f"Assistant API Error: {str(e)}")

# 📤 Upload document endpoint
@app.route("/upload-document", methods=["POST"])
def upload_document():
    # 0) remove any prior upload
    for existing in os.listdir(DOCUMENT_TO_ANALYZE_PATH):
        os.remove(os.path.join(DOCUMENT_TO_ANALYZE_PATH, existing))

    # 1) now save the new file
    if 'file' not in request.files or request.files['file'].filename == '':
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    file.save(os.path.join(DOCUMENT_TO_ANALYZE_PATH, file.filename))
    return jsonify({"message": "File uploaded successfully", "filename": file.filename}), 200


# 📄 List proprietary documents
@app.route("/list-proprietary-documents", methods=["GET"])
def list_proprietary_documents():
    try:
        files = [f for f in os.listdir(PROPRIETARY_FOLDER) if f.endswith(".docx")]
        return jsonify({"documents": files})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# 📥 Upload proprietary document
@app.route("/upload-proprietary-document", methods=["POST"])
def upload_proprietary_document():
    if 'file' not in request.files or request.files['file'].filename == '':
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    file.save(os.path.join(PROPRIETARY_FOLDER, file.filename))
    return jsonify({"message": "File uploaded successfully", "filename": file.filename}), 200

# 🗑️ Delete proprietary document
@app.route("/delete-proprietary-document", methods=["DELETE"])
def delete_proprietary_document():
    filename = request.json.get("filename")
    if not filename:
        return jsonify({"error": "Filename is required"}), 400

    path = os.path.join(PROPRIETARY_FOLDER, filename)
    if os.path.exists(path):
        os.remove(path)
        return jsonify({"message": "File deleted successfully"}), 200
    return jsonify({"error": "File not found"}), 404

# 📊 Analyze endpoint — sends 4 prompts to the assistant
@app.route("/analyze", methods=["POST"])
def analyze_document():
    # ensure the trace matrix is loaded (only actually ingests once)
    app.logger.debug("TraceMatrix before ingest: %d", collection.count())
    ingest_trace_matrix()
    app.logger.debug("TraceMatrix after  ingest: %d", collection.count())
    try:

        # 1. Load uploaded document
        files = [f for f in os.listdir(DOCUMENT_TO_ANALYZE_PATH) if f.endswith(".docx")]
        if not files:
            return jsonify({"error": "No uploaded document found"}), 400

        doc_path = os.path.join(DOCUMENT_TO_ANALYZE_PATH, files[0])
        document_text = read_docx(doc_path)

        # 2. Extract operational testing section
        operational_text = extract_operational_testing_section(document_text)
        if not operational_text:
            return jsonify({"error": "Operational Testing section not found"}), 400

        # 3. Save inputs for verification
        save_to_docx(document_text, os.path.join(RESULTS_FOLDER, "analyzed_document.docx"))
        save_to_docx(operational_text, os.path.join(RESULTS_FOLDER, "operational_testing_section.docx"))

        # 4. Load proprietary docs and save full input
        proprietary_full_text = ""
        for file in os.listdir(PROPRIETARY_FOLDER):
            if file.endswith(".docx"):
                proprietary_full_text += f"\n--- {file} ---\n{read_docx(os.path.join(PROPRIETARY_FOLDER, file))}\n"

        save_to_docx(proprietary_full_text, os.path.join(RESULTS_FOLDER, "proprietary_documents_input.docx"))

        # 5. Prepare four prompt sections
        prompt_part1 = f"""
            You are an expert Compliance and Validation Analyst specializing in regulatory compliance, risk assessment, and test plan validation.
            Your task is to analyze a regulatory document ("document_to_analyze") against several reference documents located in the "proprietary_documents" folder.
            Treat this as a new, independent request. 
            Reference documents include:
            - GAMP5 - Validated IT Systems Info
            - Guidance-Computer-Software-Assurance
            - Test Plan Template
            - Trace Matrix v8
            - Risk Levels

            1. Compliance Findings (GAMP5 & CSA Standards)

            Compare the document_to_analyze to the GAMP5 and CSA guidance documents.
            For each compliance issue found:
                - Issue: The non-compliant text or policy.
                - Section: Where in the document the issue appears.
                - Regulatory Reference: The specific GAMP5 or CSA principle violated.
                - Correction: A recommended fix.
            Notes:
                - Focus on compliance-related gaps — skip document approval sections and IQ/OQ scripts.
                - Provide precise reasoning.
            Output Format:
            Compliance Findings (GAMP5 & CSA Standards)
                Issue: [Description]
                Section: [Section title and reference]
                Regulatory Reference: [Reference]
                Correction: [Fix]

            {document_text}
        """

        result1, tokens1, cost1, time1 = run_prompt_with_assistant(
            assistant_id="asst_ByTe0UXgoT8EYqWwU4XBNCvH",
            user_input=prompt_part1
        )
        print(f"✅ Prompt 1 done. Tokens: {tokens1}, Cost: ${cost1:.4f}, Time: {time1:.2f}s")

        prompt_part2 = f"""
            2. Structural & Consistency Findings (Test Plan Alignment)

            Compare the structure of the document_to_analyze to the Test Plan Template and report any of the following:
            - Misaligned or missing headings
            - Incorrect section ordering
            - Structural inconsistencies

            For each structural issue:
                - Issue: Description of what's misaligned or missing.
                - Location: Section and line reference.
                - Correction: Instruction on how to align it with the Test Plan Template.

            {document_text}
        """

        result2, tokens2, cost2, time2 = run_prompt_with_assistant(
            assistant_id="asst_ByTe0UXgoT8EYqWwU4XBNCvH",
            user_input=prompt_part2
        )
        print(f"✅ Prompt 2 done. Tokens: {tokens2}, Cost: ${cost2:.4f}, Time: {time2:.2f}s")


        prompt_part3 = f"""
            3. System Name Consistency Check

            Identify the official system name from the first page of the document.
            Then scan the entire document for any other names or variants used for the system.
            Report names that do not match the official system name.
            For each inconsistent reference:
                - Incorrect Name Used: Exact name used.
                - Sentence: Full sentence containing the incorrect name.
                - Correction: Suggest replacement with the correct system name.

            {document_text}
        """

        result3, tokens3, cost3, time3 = run_prompt_with_assistant(
            assistant_id="asst_ByTe0UXgoT8EYqWwU4XBNCvH",
            user_input=prompt_part3
        )
        print(f"✅ Prompt 3 done. Tokens: {tokens3}, Cost: ${cost3:.4f}, Time: {time3:.2f}s")


        print("===== OPERATIONAL_TEXT =====")
        print(operational_text)
        matches = re.findall(
            r"([\s\S]*?)\s*Requirement\s*#?:?\s*([A-Za-z0-9\-]+)",
            operational_text.strip()
        )
        print("Regex found matches:", matches)

        app.logger.debug("Operational text:\n%s", operational_text)
        app.logger.debug("Regex matches: %s", matches)

        # 4a) Build a list of (change_description, chosen_req) tuples
        matches = []
        lines = operational_text.splitlines()
        i = 0
        while i < len(lines) - 1:
            title_line = lines[i].strip()
            id_line    = lines[i+1].strip()

            if title_line.endswith(":") and re.match(r"^(?:BR|FR|UR-REG|FS-REG)\s*[\d\.]+", id_line):
                change_description = title_line.rstrip(":")
                m = re.match(r"^([A-Z0-9\-]+\s*\d+(?:\.\d+)*)", id_line)
                chosen_req = m.group(1) if m else id_line
                matches.append((change_description, chosen_req))
                i += 2
            else:
                i += 1

        app.logger.debug("Built %d change/ID pairs for RAG", len(matches))

        # 4b) For each pair, retrieve top‑5 IDs and format the section
        result4_sections = []
        for change_description, chosen_req in matches:
            top_pairs = retrieve_relevant_trace_requirements(change_description, top_k=5)

            section_lines = [
                f"Change Description: {change_description}",
                f"Chosen Requirement: {chosen_req}",
                "AI Chosen Impacted Requirements:"
            ]
            # ← insert your req_id/req_text loop here:
            for req_id, req_text in top_pairs:
                section_lines.append(f"    {req_id}")
                section_lines.append(f"        {req_text}")

            result4_sections.append("\n".join(section_lines))

        # 4c) Join all sections with dividers
        result4 = "\n\n---\n\n".join(result4_sections)
        app.logger.debug("Prompt 4 sections built: %d", len(result4_sections))
        final_result = "\n\n".join([result1, result2, result3, result4])

        total_tokens = tokens1 + tokens2 + tokens3  # note: no tokens4 now
        total_cost   = cost1   + cost2   + cost3
        total_time   = time1   + time2   + time3

        # 7. Save result
        save_to_docx(final_result,
                     os.path.join(RESULTS_FOLDER, "formatted_analysis.docx"))

        return jsonify({
            "result": final_result,
            "saved_path": os.path.join(RESULTS_FOLDER, "formatted_analysis.docx"),
            "tokens_used": total_tokens,
            "cost": total_cost,
            "elapsed_time": total_time
        })

    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        return jsonify({"error": str(e)}), 500

# 📥 Download final formatted results
@app.route("/download-results", methods=["GET"])
def download_results():
    path = os.path.join(RESULTS_FOLDER, "formatted_analysis.docx")
    if not os.path.exists(path):
        return jsonify({"error": "Formatted analysis file not found"}), 404
    return send_file(path, as_attachment=True)

# 🚀 Run server
if __name__ == "__main__":
    # bind to 0.0.0.0 so Docker port‑forwarding actually works
    app.run(host="0.0.0.0", port=5000, debug=True)
